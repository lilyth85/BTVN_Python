from flask import Flask, render_template, Response, request, redirect
from pymysql import connect, cursors, Error
from datetime import datetime  
from docx import Document
from docx.shared import Inches
#from app import app


config = {
    'host': 'localhost',
    'user': 'root',
    'password': '',
    'database': 'blog',
    }
cnx = connect(**config)

app = Flask(__name__)

@app.route('/', methods=["GET"])
def render_form():
    cur = cnx.cursor()
    sql = "select * from blog_content"
    cur.execute(sql)
    #cnx.close()
    return render_template("index.html",rows=cur)


@app.route('/about')
def about():
    cur = cnx.cursor()
    sql = "select * from blog_content"
    cur.execute(sql)
    #cnx.close()
    return render_template("about.html")



@app.route('/', methods=["POST"])
def new_post():
    cur = cnx.cursor()
    title = request.form["title"]
    content = request.form["content"]
    sql = "INSERT INTO blog_content (title, content, date) VALUES (%s, %s,%s)"
    value = (title,content,datetime.now())
    try:
        cur.execute(sql,value)
        cnx.commit()
    except:
        cnx.rollback()
    #cnx.close()
    return redirect("/", code=302)

@app.route('/editpost/<id>', methods=["GET"])
def render_form_edit(id):
    cur = cnx.cursor()
    sql = "select * from blog_content where id="+str(id)
    cur.execute(sql)
    for row in cur:
         r0 = row[0]
         r1 = row[1]
         r2 = row[2]


    #cnx.close()
    return render_template("editpost.html",r0=r0,r1=r1,r2=r2)



@app.route('/editpost/<id>', methods=["POST"])
def edit_post(id):
    cur = cnx.cursor()
    title = request.form["title"]
    content = request.form["content"]
    sql = f" UPDATE blog_content SET title='{title}',content='{content}' where id = {id}"
    try:
        cur.execute(sql)
        cnx.commit()
    except:
        cnx.rollback()
    #cnx.close()
    return redirect("/editpost/"+id, code=302)


@app.route('/letter', methods=["GET"])
def xport_file():
    #cnx.close()
    return render_template("letter.html")

@app.route('/letter', methods=["POST"])
def get_text():
    name = request.form["name"]
    reason = request.form["reason"]
    document = Document()
    document.add_heading(name, 0)
    document.add_paragraph(reason)
    document.save('static/demo.docx')
    return redirect("/letter", code=302)


    

if __name__ == "__main__":
    app.run(debug = True)